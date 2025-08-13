import streamlit as st
import pandas as pd
import PyPDF2
from docx import Document
from docx.shared import Inches
import json
import os
from openai import OpenAI
import zipfile
import io
import numpy as np
from typing import Dict, List, Tuple, Any

def get_llm_response(transcript_text: str, features: List[str], model: str) -> Dict[str, int]:
    """Gets the LLM response for the transcript with binary values."""
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=st.secrets["OPENROUTER_API_KEY"],
    )
    
    features_list = "\n".join([f"- {feature}" for feature in features])
    
    prompt = f"""
Given the following clinical transcript and list of features, determine whether each feature has been explicitly addressed in the transcript.

A feature is considered **"Addressed"** if the transcript provides any relevant information affirming, denying, or otherwise commenting on the feature — even if only partially or indirectly (e.g., "no difficulty swallowing" would Address "dysphagia").

A feature is considered **"Not Addressed"** if it is not mentioned or inferable at all.

Respond in JSON format, where each key is a feature and each value is either 1 (if addressed) or 0 (if not addressed).

Transcript:
{transcript_text}

Features:
{features_list}

Return only the JSON object with no additional text.
    """

    completion = client.chat.completions.create(
        extra_headers={
            "HTTP-Referer": "https://streamlit.io", 
            "X-Title": "Transcript Feature Checker",
        },
        model=model,
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        response_format={"type": "json_object"}
    )
    return json.loads(completion.choices[0].message.content)

def extract_text_from_pdf(pdf_file) -> str:
    """Extracts text from a PDF file."""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_txt(txt_file) -> str:
    """Extracts text from a TXT file."""
    # Read the file content as string
    content = txt_file.read()
    # Handle both bytes and string content
    if isinstance(content, bytes):
        return content.decode('utf-8')
    return content

def extract_text_from_docx(docx_file) -> str:
    """Extracts text from a DOCX file."""
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_file(uploaded_file) -> str:
    """Extracts text from uploaded file based on file type."""
    file_extension = uploaded_file.name.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == 'txt':
        return extract_text_from_txt(uploaded_file)
    elif file_extension == 'docx':
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def parse_feature_lr_matrix(uploaded_file) -> Tuple[pd.DataFrame, List[str]]:
    """Parse the uploaded feature-LR matrix file."""
    if uploaded_file.name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)
    
    # First column should be features, remaining columns are diagnostic categories
    features = df.iloc[:, 0].tolist()
    diagnostic_categories = df.columns[1:].tolist()
    
    # Set features as index for easier lookup
    df.set_index(df.columns[0], inplace=True)
    
    return df, diagnostic_categories

def validate_feature_lr_matrix(df: pd.DataFrame) -> bool:
    """Validate that the feature-LR matrix has proper structure."""
    # Check that all LR values are positive numbers
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    if len(numeric_cols) == 0:
        return False
    
    # Check for negative values
    if (df[numeric_cols] < 0).any().any():
        return False
    
    return True

def validate_prior_probabilities(priors: Dict[str, float]) -> bool:
    """Validate prior probabilities.
    
    For single diagnostic category: Any positive value is valid (represents pretest probability)
    For multiple categories: Must sum to 1.0 (competing diagnoses)
    """
    if len(priors) == 1:
        # Single category: just needs to be positive (will be used as pretest probability)
        return all(prob > 0 for prob in priors.values())
    else:
        # Multiple categories: must sum to 1.0
        total = sum(priors.values())
        return abs(total - 1.0) < 0.001

def calculate_entropy(probabilities: np.ndarray) -> float:
    """Calculate entropy using log base 2: -Σ(p * log₂(p))"""
    # Handle zero probabilities to avoid log(0)
    probabilities = np.maximum(probabilities, 1e-10)
    return -np.sum(probabilities * np.log2(probabilities))

def calculate_kl_divergence(posterior: np.ndarray, prior: np.ndarray) -> float:
    """Calculate KL divergence using log base 2: Σ(q * log₂(q/p))"""
    # Handle zero probabilities to avoid log(0)
    posterior = np.maximum(posterior, 1e-10)
    prior = np.maximum(prior, 1e-10)
    return np.sum(posterior * np.log2(posterior / prior))

def perform_belief_update(prior_probs: np.ndarray, likelihood_ratios: np.ndarray) -> Tuple[np.ndarray, float]:
    """
    Perform Bayesian belief updating with likelihood ratios using normalized odds approach.
    
    This implementation matches the spreadsheet methodology recommended by clinical experts
    for more conservative probability estimates that maintain diagnostic uncertainty longer.
    
    SINGLE CATEGORY: Uses binary odds form (disease vs no disease)
    - Converts prior to odds: odds = prior / (1 - prior)
    - Applies LR: posterior_odds = prior_odds × LR
    - Converts back: probability = posterior_odds / (1 + posterior_odds)
    
    MULTIPLE CATEGORIES: Uses normalized odds approach
    - Converts each prior to odds individually
    - Applies LR to each: posterior_odds = prior_odds × LR
    - Converts back to probabilities: prob = posterior_odds / (1 + posterior_odds)
    - Normalizes all probabilities to sum to 1.0
    
    This approach is more conservative than direct probability multiplication,
    maintaining uncertainty longer and allowing later features to have meaningful impact.
    """
    
    # Validate inputs
    if len(prior_probs) != len(likelihood_ratios):
        raise ValueError("Prior probabilities and likelihood ratios must have same length")
    
    if np.any(likelihood_ratios <= 0):
        raise ValueError("Likelihood ratios must be positive")
    
    # Handle single vs multiple categories differently
    if len(prior_probs) == 1:
        # SINGLE CATEGORY: Binary Bayesian updating
        if np.any(prior_probs <= 0):
            raise ValueError("Prior probability must be positive")
        
        prior_prob = prior_probs[0]
        lr = likelihood_ratios[0]
        
        # Convert to odds form for proper LR application
        prior_odds = prior_prob / (1 - prior_prob)
        posterior_odds = prior_odds * lr
        posterior_prob = posterior_odds / (1 + posterior_odds)
        
        return np.array([posterior_prob]), posterior_odds
        
    else:
        # MULTIPLE CATEGORIES: Use normalized odds approach (matches spreadsheet)
        prior_sum = np.sum(prior_probs)
        
        # Auto-normalize priors if they're close to 1.0 but not exact (handles rounding errors)
        if abs(prior_sum - 1.0) < 0.02:  # Allow up to 2% deviation
            prior_probs = prior_probs / prior_sum  # Normalize to sum to 1.0
        elif not np.allclose(prior_sum, 1.0, atol=1e-6):
            raise ValueError(f"Prior probabilities must sum to 1.0. Current sum: {prior_sum:.6f}")
        
        # NORMALIZED ODDS METHOD (matches spreadsheet approach)
        # This is more conservative and maintains uncertainty longer
        unnormalized_odds_results = []
        
        for i in range(len(prior_probs)):
            # Convert prior to odds
            prior_odds = prior_probs[i] / (1 - prior_probs[i])
            
            # Apply LR
            posterior_odds = prior_odds * likelihood_ratios[i]
            
            # Convert back to probability (unnormalized)
            unnormalized_prob = posterior_odds / (1 + posterior_odds)
            unnormalized_odds_results.append(unnormalized_prob)
        
        unnormalized_odds_results = np.array(unnormalized_odds_results)
        
        # Normalize so all probabilities sum to 1.0
        normalization_constant = np.sum(unnormalized_odds_results)
        
        if normalization_constant == 0:
            raise ValueError("Normalization constant is zero - check likelihood ratio values")
        
        posterior_probs = unnormalized_odds_results / normalization_constant
        
        return posterior_probs, normalization_constant

def process_transcript_features(transcript_text: str, feature_lr_df: pd.DataFrame,
                              prior_probs: Dict[str, float], model: str) -> Dict[str, Any]:
    """Process transcript features sequentially and calculate belief updates."""
    
    # Get feature assessment from LLM
    features = feature_lr_df.index.tolist()
    assessment_results = get_llm_response(transcript_text, features, model)
    
    # Filter to only addressed features (value = 1)
    addressed_features = [feature for feature, value in assessment_results.items() if value == 1]
    
    # Initialize tracking variables
    diagnostic_categories = feature_lr_df.columns.tolist()
    current_probs = np.array([prior_probs[cat] for cat in diagnostic_categories])
    
    # Track calculations for each step
    calculation_steps = []
    cumulative_kl = 0.0
    initial_entropy = calculate_entropy(current_probs)
    
    # Process each addressed feature sequentially
    for i, feature in enumerate(addressed_features):
        if feature in feature_lr_df.index:
            # Get likelihood ratios for this feature
            lr_values = feature_lr_df.loc[feature].values
            
            # Calculate entropy before update
            entropy_before = calculate_entropy(current_probs)
            
            # Perform belief update
            new_probs, norm_constant = perform_belief_update(current_probs, lr_values)
            
            # Calculate entropy after update
            entropy_after = calculate_entropy(new_probs)
            
            # Calculate entropy reduction
            entropy_reduction = entropy_before - entropy_after
            
            # Calculate KL divergence for this step
            kl_divergence = calculate_kl_divergence(new_probs, current_probs)
            cumulative_kl += kl_divergence
            
            # Store step information with detailed calculation breakdown
            step_info = {
                'step': i + 1,
                'feature': feature,
                'prior_probs': current_probs.copy(),
                'likelihood_ratios': lr_values.copy(),
                'unnormalized_posterior': current_probs * lr_values,  # For debugging
                'posterior_probs': new_probs.copy(),
                'entropy_before': entropy_before,
                'entropy_after': entropy_after,
                'entropy_reduction': entropy_reduction,
                'kl_divergence': kl_divergence,
                'normalization_constant': norm_constant,
                'manual_calculation_check': {
                    'prior_x_lr': (current_probs * lr_values).tolist(),
                    'sum_for_normalization': norm_constant,
                    'final_probabilities': new_probs.tolist()
                }
            }
            calculation_steps.append(step_info)
            
            # Update current probabilities for next iteration
            current_probs = new_probs
    
    # Calculate final statistics
    final_entropy = calculate_entropy(current_probs)
    total_entropy_reduction = initial_entropy - final_entropy
    
    return {
        'assessment_results': assessment_results,
        'addressed_features': addressed_features,
        'calculation_steps': calculation_steps,
        'diagnostic_categories': diagnostic_categories,
        'initial_entropy': initial_entropy,
        'final_entropy': final_entropy,
        'total_entropy_reduction': total_entropy_reduction,
        'cumulative_kl_divergence': cumulative_kl,
        'final_probabilities': dict(zip(diagnostic_categories, current_probs))
    }

def create_comprehensive_docx(transcript_name: str, all_results: Dict[str, Dict[str, Any]],
                             all_lr_matrices: Dict[str, Dict], model: str) -> str:
    """Creates a comprehensive DOCX file with results from all LR matrices for a single transcript."""
    document = Document()
    document.add_heading(f'Comprehensive Transcript Analysis Report for {transcript_name}', 0)
    
    # Add model information and overview
    document.add_paragraph(f'Analysis performed using: {model}')
    document.add_paragraph(f'Number of LR matrices analyzed: {len(all_results)}')
    document.add_paragraph('')
    
    # Executive Summary comparing all matrices
    document.add_heading('Executive Summary - Comparison Across LR Matrices', level=1)
    
    summary_table = document.add_table(rows=1, cols=5)
    summary_table.style = 'Table Grid'
    
    # Headers
    summary_headers = summary_table.rows[0].cells
    summary_headers[0].text = 'LR Matrix'
    summary_headers[1].text = 'Features Addressed'
    summary_headers[2].text = 'Entropy Reduction'
    summary_headers[3].text = 'KL Divergence'
    summary_headers[4].text = 'Top Diagnosis'
    
    # Add summary data for each matrix
    for lr_filename, results in all_results.items():
        row_cells = summary_table.add_row().cells
        row_cells[0].text = lr_filename
        row_cells[1].text = f"{len(results['addressed_features'])}/{len(results['assessment_results'])}"
        row_cells[2].text = f"{results['total_entropy_reduction']:.3f}"
        row_cells[3].text = f"{results['cumulative_kl_divergence']:.3f}"
        
        # Find top diagnosis
        top_diagnosis = max(results['final_probabilities'].items(), key=lambda x: x[1])
        row_cells[4].text = f"{top_diagnosis[0][:20]}... ({top_diagnosis[1]:.3f})"
    
    document.add_paragraph('')
    
    # Detailed analysis for each LR matrix
    for lr_filename, results in all_results.items():
        feature_lr_df = all_lr_matrices[lr_filename]['df']
        
        document.add_heading(f'Detailed Analysis: {lr_filename}', level=1)
        
        # 1. Feature Assessment Summary for this matrix
        document.add_heading('Feature Assessment Summary', level=2)
        total_features = len(results['assessment_results'])
        addressed_count = len(results['addressed_features'])
        not_addressed_count = total_features - addressed_count
        
        matrix_summary_table = document.add_table(rows=4, cols=2)
        matrix_summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Total Features Assessed', str(total_features)),
            ('Features Addressed', str(addressed_count)),
            ('Features Not Addressed', str(not_addressed_count)),
            ('Percentage Addressed', f'{(addressed_count/total_features)*100:.1f}%')
        ]
        
        for i, (label, value) in enumerate(summary_data):
            matrix_summary_table.cell(i, 0).text = label
            matrix_summary_table.cell(i, 1).text = value
        
        document.add_paragraph('')
        
        # 2. Addressed Features with Likelihood Ratios for this matrix
        if results['addressed_features']:
            document.add_heading('Addressed Features with Likelihood Ratios', level=2)
            
            lr_table = document.add_table(rows=1, cols=len(results['diagnostic_categories']) + 1)
            lr_table.style = 'Table Grid'
            
            # Headers
            lr_headers = lr_table.rows[0].cells
            lr_headers[0].text = 'Feature'
            for i, category in enumerate(results['diagnostic_categories']):
                lr_headers[i + 1].text = category[:15] + '...' if len(category) > 15 else category
            
            # Add addressed features with their LR values
            for feature in results['addressed_features']:
                if feature in feature_lr_df.index:
                    row_cells = lr_table.add_row().cells
                    row_cells[0].text = feature
                    lr_values = feature_lr_df.loc[feature].values
                    for i, lr_value in enumerate(lr_values):
                        row_cells[i + 1].text = f'{lr_value:.2f}'
            
            document.add_paragraph('')
        
        # 3. Belief Updating Sequence for this matrix
        if results['calculation_steps']:
            document.add_heading('Belief Updating Sequence', level=2)
            
            sequence_table = document.add_table(rows=1, cols=6)
            sequence_table.style = 'Table Grid'
            
            # Headers
            seq_headers = sequence_table.rows[0].cells
            seq_headers[0].text = 'Step'
            seq_headers[1].text = 'Feature'
            seq_headers[2].text = 'Entropy Before'
            seq_headers[3].text = 'Entropy After'
            seq_headers[4].text = 'Entropy Reduction'
            seq_headers[5].text = 'KL Divergence'
            
            # Add calculation steps
            for step in results['calculation_steps']:
                row_cells = sequence_table.add_row().cells
                row_cells[0].text = str(step['step'])
                row_cells[1].text = step['feature'][:20] + '...' if len(step['feature']) > 20 else step['feature']
                row_cells[2].text = f'{step["entropy_before"]:.3f}'
                row_cells[3].text = f'{step["entropy_after"]:.3f}'
                row_cells[4].text = f'{step["entropy_reduction"]:.3f}'
                row_cells[5].text = f'{step["kl_divergence"]:.3f}'
            
            document.add_paragraph('')
        
        # 4. Final Diagnostic Probabilities for this matrix
        document.add_heading('Final Diagnostic Probabilities', level=2)
        
        prob_table = document.add_table(rows=1, cols=2)
        prob_table.style = 'Table Grid'
        
        prob_headers = prob_table.rows[0].cells
        prob_headers[0].text = 'Diagnostic Category'
        prob_headers[1].text = 'Final Probability'
        
        for category, probability in results['final_probabilities'].items():
            row_cells = prob_table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = f'{probability:.4f}'
        
        document.add_paragraph('')
        
        # 5. Summary Statistics for this matrix
        document.add_heading('Summary Statistics', level=2)
        
        stats_table = document.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Initial Entropy', f'{results["initial_entropy"]:.3f} bits'),
            ('Final Entropy', f'{results["final_entropy"]:.3f} bits'),
            ('Total Entropy Reduction', f'{results["total_entropy_reduction"]:.3f} bits'),
            ('Cumulative KL Divergence', f'{results["cumulative_kl_divergence"]:.3f} bits'),
            ('Number of Belief Updates', str(len(results['calculation_steps'])))
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = value
        
        document.add_paragraph('')
        document.add_page_break()  # Separate each matrix analysis
    
    # Save the document
    file_path = f"comprehensive_analysis_{transcript_name}.docx"
    document.save(file_path)
    return file_path

def check_password():
    """Returns `True` if the user had a correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["PASSWORD"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 Password incorrect")
        return False
    else:
        # Password correct.
        return True

if check_password():
    st.title('Enhanced Transcript Feature Analysis with Belief Updating')

    st.sidebar.header('Model Selection')
    model_options = [
        "google/gemini-2.5-flash",
        "google/gemini-2.5-pro",
        "openai/gpt-4o-mini",
        "openai/gpt-4o",
        "openai/gpt-4.1-mini",        
        "openai/gpt-4.1",
        "openai/o3-mini",
        "openai/gpt-5",
        "anthropic/claude-3.7-sonnet",
        "anthropic/claude-sonnet-4"
    ]
    selected_model = st.sidebar.selectbox("Choose a model", model_options, index=4)

    # Step 1: Upload Feature-LR Matrices
    st.header('Step 1: Upload Feature-Likelihood Ratio Matrices')
    st.write("Upload one or more CSV or Excel files with different LR configurations (tiers/groupings). Each file should have features in the first column and diagnostic categories as column headers with their likelihood ratios.")
    
    uploaded_lr_files = st.file_uploader(
        "Upload Feature-LR Matrices",
        type=['csv', 'xlsx'],
        accept_multiple_files=True,
        help="First column: Features, Remaining columns: Diagnostic categories with LR values"
    )

    if uploaded_lr_files:
        lr_matrices = {}
        all_valid = True
        
        # Process each uploaded LR matrix
        for lr_file in uploaded_lr_files:
            try:
                feature_lr_df, diagnostic_categories = parse_feature_lr_matrix(lr_file)
                
                if validate_feature_lr_matrix(feature_lr_df):
                    lr_matrices[lr_file.name] = {
                        'df': feature_lr_df,
                        'categories': diagnostic_categories
                    }
                    st.success(f"✅ {lr_file.name}: Loaded {len(feature_lr_df)} features and {len(diagnostic_categories)} diagnostic categories")
                else:
                    st.error(f"❌ {lr_file.name}: Invalid LR matrix. Please ensure all LR values are positive numbers.")
                    all_valid = False
            except Exception as e:
                st.error(f"❌ Error processing {lr_file.name}: {str(e)}")
                all_valid = False
        
        if all_valid and lr_matrices:
            # Display previews of all matrices
            st.write("**Preview of Feature-LR Matrices:**")
            for filename, data in lr_matrices.items():
                with st.expander(f"Preview: {filename}"):
                    st.dataframe(data['df'].head())
            
            # Step 2: Set Prior Probabilities for each matrix
            st.header('Step 2: Set Prior Probabilities for Each LR Matrix')
            st.write("Set prior probabilities for each uploaded LR matrix. This allows you to compare different LR configurations with different prior assumptions.")
            
            all_prior_probs = {}
            all_priors_valid = True
            
            for filename, data in lr_matrices.items():
                st.subheader(f"Prior Probabilities for: {filename}")
                diagnostic_categories = data['categories']
                
                if len(diagnostic_categories) == 1:
                    st.write("Enter the pretest probability for the diagnostic category (any positive value, e.g., 0.35 for 35%).")
                else:
                    st.write("Enter the prior probability for each diagnostic category. Values must sum to 1.0.")
                
                # Create input fields for prior probabilities
                prior_probs = {}
                cols = st.columns(min(3, len(diagnostic_categories)))
                
                for i, category in enumerate(diagnostic_categories):
                    col_idx = i % len(cols)
                    with cols[col_idx]:
                        # Use filename in key to make it unique across matrices
                        safe_filename = filename.replace('.', '_').replace(' ', '_')
                        prior_probs[category] = st.number_input(
                            f"{category[:20]}{'...' if len(category) > 20 else ''}",
                            min_value=0.0,
                            max_value=1.0,
                            value=1.0/len(diagnostic_categories),  # Default to uniform distribution
                            step=0.001,
                            format="%.3f",
                            key=f"prior_{safe_filename}_{i}"
                        )
                
                # Validate prior probabilities for this matrix
                if validate_prior_probabilities(prior_probs):
                    if len(diagnostic_categories) == 1:
                        st.success(f"✅ Pretest probability set to {list(prior_probs.values())[0]:.3f} ({list(prior_probs.values())[0]*100:.1f}%)")
                    else:
                        st.success(f"✅ Prior probabilities sum to {sum(prior_probs.values()):.3f}")
                    all_prior_probs[filename] = prior_probs
                else:
                    if len(diagnostic_categories) == 1:
                        st.error(f"❌ Pretest probability must be positive. Current value: {list(prior_probs.values())[0]:.3f}")
                    else:
                        st.error(f"❌ Prior probabilities must sum to 1.0. Current sum: {sum(prior_probs.values()):.3f}")
                    all_priors_valid = False
                
                st.write("---")  # Separator between matrices
            
            # Only proceed if all prior probabilities are valid
            if all_priors_valid and len(all_prior_probs) == len(lr_matrices):
                    
                    # Step 3: Upload Transcripts
                    st.header('Step 3: Upload Transcripts')
                    uploaded_transcripts = st.file_uploader(
                        "Upload student transcripts (PDF, TXT, or DOCX files)",
                        type=['pdf', 'txt', 'docx'],
                        accept_multiple_files=True
                    )

                    if uploaded_transcripts:
                        if st.button('Generate Enhanced Analysis Reports'):
                            with st.spinner('Processing transcripts and performing belief updating calculations...'):
                                zip_buffer = io.BytesIO()
                                docx_files = []
                                
                                with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                                    # Process each transcript against all LR matrices
                                    for transcript_file in uploaded_transcripts:
                                        try:
                                            # Extract text from file (PDF, TXT, or DOCX)
                                            transcript_text = extract_text_from_file(transcript_file)
                                            transcript_name = os.path.splitext(transcript_file.name)[0]
                                            
                                            # Store results from all LR matrices for this transcript
                                            all_results_for_transcript = {}
                                            
                                            # Process transcript with each LR matrix
                                            for lr_filename, lr_data in lr_matrices.items():
                                                feature_lr_df = lr_data['df']
                                                prior_probs = all_prior_probs[lr_filename]
                                                
                                                # Process transcript with belief updating
                                                results = process_transcript_features(
                                                    transcript_text, feature_lr_df, prior_probs, selected_model
                                                )
                                                
                                                # Store results for comprehensive report
                                                all_results_for_transcript[lr_filename] = results
                                                
                                                # Display summary for this transcript-matrix combination
                                                st.success(f"✅ Processed {transcript_file.name} with {lr_filename}")
                                                st.write(f"- Features addressed: {len(results['addressed_features'])}/{len(results['assessment_results'])}")
                                                st.write(f"- Total entropy reduction: {results['total_entropy_reduction']:.3f} bits")
                                                st.write(f"- Cumulative KL divergence: {results['cumulative_kl_divergence']:.3f} bits")
                                            
                                            # Generate comprehensive DOCX report for this transcript
                                            comprehensive_docx_path = create_comprehensive_docx(
                                                transcript_name, all_results_for_transcript, lr_matrices, selected_model
                                            )
                                            docx_files.append(comprehensive_docx_path)
                                            
                                            # Add comprehensive report to ZIP
                                            zip_file.write(comprehensive_docx_path)
                                            
                                            st.info(f"📄 Generated comprehensive report: {os.path.basename(comprehensive_docx_path)}")
                                            st.write("---")  # Separator between transcripts
                                            
                                        except Exception as e:
                                            st.error(f"❌ Error processing {transcript_file.name}: {str(e)}")
                                
                                # Clean up temporary files
                                for file_path in docx_files:
                                    if os.path.exists(file_path):
                                        os.remove(file_path)

                                # Provide download button
                                st.download_button(
                                    label="📥 Download All Enhanced Reports as ZIP",
                                    data=zip_buffer.getvalue(),
                                    file_name="enhanced_transcript_analysis.zip",
                                    mime="application/zip"
                                )
            else:
                st.warning("⚠️ Please ensure all prior probabilities are valid before proceeding.")
        else:
            st.warning("⚠️ Please upload valid Feature-LR matrices before proceeding.")
    else:
        st.info("👆 Please upload one or more Feature-LR matrices to begin.")
